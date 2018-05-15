from django.shortcuts import render
from .models import *
from .forms import *
import datetime, hashlib
from django.contrib.auth.decorators import permission_required, login_required
from django.shortcuts import redirect, get_object_or_404
from django.core.urlresolvers import reverse
from auth_system.models import MyUser, Group
from django.http import HttpResponse,StreamingHttpResponse, HttpResponseRedirect
import os, cgi, json, html, zipfile, random, string, chardet, shutil
from django.views.generic.detail import DetailView
from django.utils.datastructures import MultiValueDictKeyError
from django.apps import apps
from django.core.exceptions import ObjectDoesNotExist
from django.db import transaction
from onlineTest.settings import BASE_DIR, USER_FILE_DIR
from enum import Enum, unique
import pdb
import docx
import xlwt
import channels

def updateLatestInfo(courseId, infoObject):
    # 将消息发送给教师最新消息查看页面
    channels.Group('codeweekTeacherLatestInfo-' + courseId).send(
        {'text': json.dumps({"time": infoObject.time.strftime('%Y-%m-%d %H:%M:%S'),
                             "content": infoObject.info})}
    )

@login_required
def course_list_for_student(request):
    nowtime = datetime.datetime.now()         #获取当前时间用来判断有关学生的课程是否在进行
    nowCourses = CodeWeekClass.objects.filter(
        students=request.user).filter(  #过滤获取学生正在进行的课程
        begin_time__lte=nowtime).filter(
        end_time__gte=nowtime
    )
    previousCourses = CodeWeekClass.objects.filter( #过滤获取已经结束的课程
        students=request.user).filter(
        end_time__lt=nowtime
    )
    futureCourses = CodeWeekClass.objects.filter(   #过滤获取还未开始的课程
        students=request.user).filter(
        begin_time__gt=nowtime
    )
    return render(request, 'code_week/course_list.html',
                  {'nowCourses' : nowCourses, 'previousCourses' : previousCourses, 'futureCourses' : futureCourses, 'position': 'code_week_list'})

@login_required
def course_list_for_teacher(request):
    if not request.user.isTeacher:
        return render(request, 'warning.html', {'info':'教师才可以操作'})
    nowtime = datetime.datetime.now()  # 获取当前时间用来判断有关老师的课程是否在进行
    nowCourses = CodeWeekClass.objects.filter(
        teacher=request.user).filter( # 过滤获取老师正在进行的课程
        begin_time__lte=nowtime).filter(
        end_time__gte=nowtime
    )
    previousCourses = CodeWeekClass.objects.filter(  # 过滤获取已经结束的课程
        teacher=request.user).filter(
        end_time__lt=nowtime)
    futureCourses = CodeWeekClass.objects.filter(    # 过滤获取还未开始的课程
        teacher=request.user).filter(
        begin_time__gt=nowtime)
    return render(request, 'code_week/courses.html',
                  {'nowCourses': nowCourses, 'previousCourses': previousCourses, 'futureCourses': futureCourses, 'position': 'code_week_teacher_course_list'})

#重点是学生的添加，这里的策略是通过学号添加，可以是学号范围，也可以是单个学号，这里设定所有合法的学号长度都是为9的
@login_required
def add_course(request):
    if not request.user.isTeacher:
        return render(request, 'warning.html', {'info':'教师才可以操作'})
    if request.method == 'POST':
        form = AddCodeWeekForm(request.POST)
        if form.is_valid():
            print("form correct")
            print(form.cleaned_data['course_name'])
            print(form.cleaned_data['begin_time'])
            print(form.cleaned_data['end_time'])
            print(form.cleaned_data['problems'])
            begin_time = form.cleaned_data['begin_time']
            end_time = form.cleaned_data['end_time']
            if end_time <= begin_time:
                return HttpResponse("时间不对")
            newCourse = CodeWeekClass(name=form.cleaned_data['course_name'],
                                      teacher=request.user,
                                      begin_time=begin_time,
                                      end_time=end_time,
                                      numberEachGroup=form.cleaned_data['maxnumber'])
            newCourse.save()
            # 添加题目
            if form.cleaned_data['problems']:
                ids = form.cleaned_data['problems'].split(',')
                try:
                    for pk in ids:
                        newCourse.problems.add(pk)
                except:
                    return HttpResponse(0)
            # 添加学生
            for stu_detail in request.POST['students'].splitlines():
                if len(stu_detail.split()) > 1:
                    id_num, username = stu_detail.split()[0], stu_detail.split()[1]
                    try:
                        student = MyUser.objects.get(id_num=id_num)
                    except:
                        student = MyUser(id_num=id_num, email=id_num + '@njupt.edu.cn', username=username)
                        student.set_password(id_num)
                        student.save()
                        student.groups.add(Group.objects.get(name='学生'))
                    if newCourse.students.all().filter(codeweekclassstudent__student=student).count() == 0: # 判断学生是否已经存在
                        if form.cleaned_data['maxnumber'] == "1":  # 没有分组操作，直接给每个学生加入一个只有自己一个人的组
                            newgroup = CodeWeekClassGroup(cwclass=newCourse)
                            newgroup.save()
                            newCourseStudent = CodeWeekClassStudent(codeWeekClass=newCourse, student=student,
                                                                    group=newgroup,
                                                                    isLeader=True)
                            newCourseStudent.save()
                        else:
                            newCourseStudent = CodeWeekClassStudent(codeWeekClass=newCourse, student=student)
                            newCourseStudent.save()
                else:
                    try:
                        student = MyUser.objects.get(id_num=stu_detail)
                        if newCourse.students.all().filter(codeweekclassstudent__student=student).count() == 0:
                            if form.cleaned_data['maxnumber'] == "1":  # 没有分组操作，直接给每个学生加入一个只有自己一个人的组
                                newgroup = CodeWeekClassGroup(cwclass=newCourse)
                                newgroup.save()
                                newCourseStudent = CodeWeekClassStudent(codeWeekClass=newCourse, student=student,
                                                                        group=newgroup,
                                                                        isLeader=True)
                                newCourseStudent.save()
                            else:
                                newCourseStudent = CodeWeekClassStudent(codeWeekClass=newCourse, student=student)
                                newCourseStudent.save()

                    except:
                        pass
            return redirect(reverse('view_course_for_teacher', args=[newCourse.id,]))
           # return view_course(request, newCourse.id)
        #    return HttpResponse(1)

        else:
            print(form.errors)
    else:
        form = AddCodeWeekForm()
        categorys = ProblemCategory.objects.all()
        return render(request, 'code_week/add_course.html', {'form': form, 'categorys': categorys})

#教师查看课程主页
@login_required
def view_course(request, courseId):
    if not request.user.isTeacher:
        return render(request, 'warning.html', {'info': '教师才可以操作'})
    course = CodeWeekClass.objects.filter(id=courseId)
    if course.count() == 0:
        return render(request, 'warning.html', {'info' : '查无此课'})
    elif course[0].teacher != request.user:
        return render(request, 'warning.html', {'info': '查无此课'})
    else:
        students = course[0].students.all()
        problems = course[0].problems.all()
        return render(request, 'code_week/course_detail.html', {'course' : course[0], 'students' : students, 'problems':problems})

#学生查看课程主页
@login_required()
def student_view_course(request, courseId):
    course = CodeWeekClass.objects.filter(id=courseId)
    if course.count() == 0:
        return render(request, 'warning.html', {'info' : '查无此课'})
    elif request.user in course[0].students.all(): # 请求的用户在课程学生列表中
        group = None
        try:
            group = CodeWeekClassStudent.objects.get(codeWeekClass=course[0], student=request.user).group
        except:
            return render(request, 'warning.html', {'info': '查询时出现问题'})
        problem = None
        if group:  # 如果还没有加入组，group为None
            problem = group.selectedProblem
        return render(request, 'code_week/student_course_detail.html', {'course' : course[0], 'problem': problem})
    else:
        return render(request, 'warning.html', {'info': '查无此课'})

# 教师提交描述文件，把文件名修改为提交的id
def newProblemFileName(fileId):
    filename = os.path.join(USER_FILE_DIR, 'upload', str(fileId))
    return filename

#增加程序设计题
@login_required
def add_sheji(request):
    if not request.user.isTeacher:
        return render(request, 'warning.html', {'info': '教师才可以操作'})
    if request.method == 'POST':  # 当提交表单时
        form = ShejiAddForm(request.POST,request.FILES)  # form 包含提交的数据
        print(form.errors)
        if form.is_valid():  # 如果提交的数据合法
            f = request.FILES.get('headImg')
            problem = form.save(user=request.user)  # 保存题目

            filename = newProblemFileName(problem.problem_id)
            fobj = open(filename,'wb')
            for chrunk in f.chunks():#可以设置分块上传
                fobj.write(chrunk)
            fobj.close()
            return redirect(reverse("sheji_detail", args=[problem.problem_id]))
    else:  # 当正常访问时
        form = ShejiAddForm()
    return render(request, 'code_week/sheji_problem_add.html', {'form': form, 'title': '新建程序设计题'})

#删除程序设计题(未完成)
@login_required
def delete_sheji(request):
    if not request.user.isTeacher:
        return render(request, 'warning.html', {'info': '教师才可以操作'})
    if request.method == 'POST':
        ids = request.POST.getlist('ids[]')
        try:
            for pk in ids:
                ShejiProblem.objects.filter(pk=pk).filter(creator=request.user).delete()
        except:
            return HttpResponse(0)
        return HttpResponse(1)
    else:
        return HttpResponse(0)

#程序设计题详细视图(未完成)
class ShejiProblemDetailView(DetailView):
    model = ShejiProblem
    template_name = 'code_week/sheji_problem_detail.html'
    context_object_name = 'problem'
    def get_context_data(self, **kwargs):
            context = super( ShejiProblemDetailView, self).get_context_data(**kwargs)
            context['title'] = '程序设计题“' + self.object.title + '”的详细信息'
            return context

#更新程序设计题
@login_required
def update_sheji(request, id):
    if not request.user.isTeacher:
        return render(request, 'warning.html', {'info': '教师才可以操作'})
    problem = get_object_or_404(ShejiProblem, pk=id)
    if problem.creator != request.user:
        return render(request, 'warning.html', {'info' : "您无法修改不是您创建的题目"})
    initial = {'title': problem.title,
               'editorText':problem.editorText,
               'category':problem.category,
               }  # 生成表单的初始化数据
    if request.method == "POST":  # 当提交表单时
        form = ShejiUpdateForm(request.POST,request.FILES)
        if form.is_valid():
            form.save(user=request.user, problemid=id)
            f = request.FILES.get('headImg')
            if f: # 重新上传了文件
                fobj = open(newProblemFileName(id), 'wb')
                for chrunk in f.chunks():
                    fobj.write(chrunk)
                fobj.close()
            return redirect(reverse("sheji_detail", args=[id]))
    return render(request, 'code_week/sheji_problem_update.html', {'form': ShejiUpdateForm(initial=initial), 'problem': problem})

# 程序设计题列表
@login_required()
def list_sheji(request):
    categorys = ProblemCategory.objects.all()
    return render(request, 'code_week/sheji_problem_list.html', context={'categorys': categorys, 'title': '程序设计题题库', 'position': 'sheji_list'})

@login_required()
def get_json_sheji(request):
    json_data = {}
    recodes = []
    offset = int(request.GET['offset'])
    limit = int(request.GET['limit'])
    category = int(request.GET['category'])

    if category != 0:
        problems = ShejiProblem.objects.filter(category=category)
    else:
        problems = ShejiProblem.objects
    try:
        problems = problems.filter(title__icontains=request.GET['search'])
    except:
        pass
    try:
        sort = request.GET['sort']
    except MultiValueDictKeyError:
        sort = 'pk'
    json_data['total'] = problems.count()
    if request.GET['order'] == 'desc':
        sort = '-' + sort
    for problem in problems.all().order_by(sort)[offset:offset + limit]:
        title = cgi.escape(problem.title)
        recode = {'pk': problem.pk, 'title': title, 'category': str(problem.category),
                  'update_date': problem.update_date.strftime('%Y-%m-%d %H:%M:%S'), 'creator': str(problem.creator),
                  'id': problem.pk}
        recodes.append(recode)
    json_data['rows'] = recodes
    return HttpResponse(json.dumps(json_data))

#用于学生查看自己课程中的所有题目
@login_required()
def get_problem_student(request, id):
    json_data = {}
    recodes = []
    course = get_object_or_404(CodeWeekClass, id=id)
    json_data['total'] = course.problems.all().count()
    problems = course.problems.all()
    i = 1
    for problem in problems.all():
        title = html.escape(problem.title)
        outline = ""
        for aline in problem.editorText.splitlines():
            outline += aline
            outline += '<br/>'
        recode = {'pk': problem.pk, 'title': title, 'category': str(problem.category),
                 'outline' : outline,
                  'id': i}
        recodes.append(recode)
        i = i + 1
    json_data['rows'] = recodes
    return HttpResponse(json.dumps(json_data))


def encodeFilename(filename):
    """
    :param filename:
    :return: 将中文的文件名编码成为Content-Disposition中能够被浏览器识别的UTF-8格式的文件名
    例如文件名为"张柯.doc",需要设置为'''attachment;filename*=UTF-8''%e5%bc%a0%e6%9f%af.doc'''
    然而str.encode("utf-8")返回的结果为b'\xe5\xbc\xa0\xe6\x9f\xaf.doc'需要转换一下
    这个函数就是将含有中文的文件名转换一下
    """
    originStr = str(filename.encode("utf-8"))
    print(originStr)
    originStr = originStr.replace("\\x","%")
    print(originStr)
    returnstr = originStr[2:len(originStr)-1]
    return returnstr

# 老师下载描述文件
@login_required
def teacher_download(request, problemId):
    if not request.user.isTeacher:
        return render(request, 'warning.html', {'info': '教师才可以操作'})
    problem = None
    try:
        problem = ShejiProblem.objects.get(problem_id=problemId)
    except:
        return render(request, 'warning.html', {'info': '找不到描述文件'})
    if problem:
        filename = newProblemFileName(problemId)
        def file_iterator(file_name, chunk_size=512):
            with open(file_name, 'rb') as f:
                while True:
                    c = f.read(chunk_size)
                    if c:
                        yield c
                    else:
                        break

        response = StreamingHttpResponse(file_iterator(filename))
        response['Content-Type'] = problem.content_type
        response['Content-Disposition'] = '''attachment;filename*= UTF-8''{0}'''.format(encodeFilename(problem.filename))
        return response
    else:
        return render(request, 'warning.html', {'info': '没有此文件'})

# 实现描述文件的下载功能
@login_required()
def download(request, courseId):
    student = None
    try:
        student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=courseId)
    except:
        return render(request, 'warning.html', {'info': '找不到描述文件'})
    if student: #能够匹配到学生
        if not student.group:
            return render(request, 'warning.html', {'info': '您还没有加入小组或者成为组长'})
        if not student.group.selectedProblem:
            return render(request, 'warning.html', {'info': '您所在的小组还没有选题'})
        file = student.group.selectedProblem
        filename = newProblemFileName(file.problem_id)

        def file_iterator(file_name, chunk_size=512):
            with open(file_name, 'rb') as f:
                while True:
                    c = f.read(chunk_size)
                    if c:
                        yield c
                    else:
                        break

        response = StreamingHttpResponse(file_iterator(filename))
        response['Content-Type'] = file.content_type
        response['Content-Disposition'] = '''attachment;filename*= UTF-8''{0}'''.format(encodeFilename(file.filename))
        return response
    else:
        return render(request, 'warning.html' ,{'info' : '没有此文件'})

# 用于学生课程界面有关分组信息的初始化操作
@login_required()
def student_info(request, courseId):
    """
    :param request:
    :param courseId: 正则匹配的课程id
    打包传输现在的学生分组情况用来初始前端显示
    """
    student = None
    course = None
    data = None
    try:
        with transaction.atomic():
            student = CodeWeekClassStudent.objects.get(codeWeekClass=courseId, student=request.user)
            course = CodeWeekClass.objects.get(id=courseId)
            if student and course:
                data = {'max': course.numberEachGroup, 'id': course.counter}
                groups = course.CodeWeekClass_group.all()
                groupsData = []
                for group in groups:
                    if group.using:
                        aGroupMemberData = []
                        aGroupData = {}
                        students = group.Group_member.all()
                        for s in students:
                            if s.isLeader:
                                aGroupData['leader'] = s.get_full_name()
                            else:
                                aGroupMemberData.append(s.get_full_name())
                        aGroupData['groupid'] = group.id
                        aGroupData['members'] = aGroupMemberData
                        if group.selectedProblem:
                            aGroupData['problem'] = group.selectedProblem.title
                        groupsData.append(aGroupData)
                    # pdb.set_trace()
                data['groups'] = groupsData
                data['name'] = student.get_full_name()
    except ObjectDoesNotExist:
        return
    return HttpResponse(json.dumps(data))

# 用于教师获取学生课程有关分组信息
@login_required()
def teacher_get_student_info(request, courseId):
    """
    :param request:
    :param courseId: 正则匹配的课程id
    打包传输现在的学生分组情况用来初始前端显示
    """
    course = None
    data = None
    try:
        with transaction.atomic():
            course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
            if course:
                data = {'id': course.counter}
                groups = course.CodeWeekClass_group.all()
                groupsData = []
                for group in groups:
                    if group.using:
                        aGroupMemberData = []
                        aGroupData = {}
                        students = group.Group_member.all()
                        for s in students:
                            if s.isLeader:
                                aGroupData['leader'] = s.get_full_name()
                            else:
                                aGroupMemberData.append(s.get_full_name())
                        aGroupData['groupid'] = group.id
                        aGroupData['members'] = aGroupMemberData
                        if group.selectedProblem:
                            aGroupData['problem'] = group.selectedProblem.title
                        groupsData.append(aGroupData)
                    # pdb.set_trace()
                data['groups'] = groupsData
                data['max'] = course.numberEachGroup
    except ObjectDoesNotExist:
        return
    return HttpResponse(json.dumps(data))

# 用于教师获取课程信息
@login_required
def teacher_course_info(request, courseId):
    course = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
    except:
        return HttpResponse(0)
    result = {'name':course.name, 'number':course.numberEachGroup,
              'begin_time': course.begin_time.strftime('%Y-%m-%d %H:%M'),
              'end_time': course.end_time.strftime('%Y-%m-%d %H:%M')}
    return HttpResponse(json.dumps(result))

# 用于教师更新课程的信息
@login_required
def teacher_update_info(request):
    if request.method == 'POST':
        # 获取需要更新信息的课程id和需要更新的内容
        courseId = None
        action = None
        try:
            courseId = int(request.POST['id'])
            action = request.POST['action']
        except:
            return HttpResponse(0)
        # 检查这个课程是否是用户的
        course = None
        try:
            course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
        except:
            return HttpResponse(0)
        if action == 'update_name': # 更新课程的名称
            name = None
            try:
                name = request.POST['name']
            except:
                return HttpResponse(0)
            try:
                with transaction.atomic():
                    course.name = name
                    course.save()
            except:
                return HttpResponse(0)
            return HttpResponse(1)
        elif action == 'update_begin_time': # 更新课程开始时间
            time = None
            try:
                timestr = request.POST['time']
                time = datetime.datetime.strptime(timestr, '%Y-%m-%d %H:%M')
            except:
                return HttpResponse(0)
            try:
                with transaction.atomic():
                    if time < course.end_time:
                        course.begin_time = time
                        course.save()
                    else:
                        raise Exception
            except:
                return HttpResponse(0)
            return HttpResponse(1)
        elif action == 'update_end_time': # 更新课程结束时间
            time = None
            try:
                timestr = request.POST['time']
                time = datetime.datetime.strptime(timestr, '%Y-%m-%d %H:%M')
            except:
                return HttpResponse(0)
            try:
                with transaction.atomic():
                    if time > course.begin_time:
                        course.end_time = time
                        course.save()
                    else:
                        raise Exception
            except:
                return HttpResponse(0)
            return HttpResponse(1)

@login_required()
# 用于教师课程界面获取已经选取的题目
def get_select_problem(request, courseId):
    course = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
    except:
        return HttpResponse(0)
    json_data = {}
    recodes = []

    problems = course.problems.all()

    json_data['total'] = problems.count()

    for problem in problems.all():
        title = html.escape(problem.title)
        recode = {'pk': problem.pk, 'title': title, 'category': str(problem.category),
                  'update_date': problem.update_date.strftime('%Y-%m-%d %H:%M:%S'), 'creator': str(problem.creator),
                  'id': problem.pk}
        recodes.append(recode)
    json_data['rows'] = recodes
    return HttpResponse(json.dumps(json_data))

@login_required
# 用于教师移除已经选好的题目
def remove_select_problem(request):
    if request.method == 'POST':
        courseId = None
        problemId = None
        try:
            courseId = request.POST['courseId']
            problemId = request.POST['problemId']
        except:
            return HttpResponse(0)
        course = None
        problem = None
        try:
            course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
            problem = ShejiProblem.objects.get(problem_id=problemId)
        except:
            return HttpResponse(0)
        try:
            with transaction.atomic():
                # 检查是否有学生已经选了这个题目
                for group in course.CodeWeekClass_group.all():
                    if group.selectedProblem == problem:
                        return HttpResponse(2)
                course.problems.remove(problem)
        except:
            return HttpResponse(0)
        return HttpResponse(1)

@login_required
# 用于教师给课程增加题目和学生
def add_problem_student(request, courseId):
    if request.method == 'POST':
        course = None
        try:
            course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
        except:
            return HttpResponse(0)
        form = UpdateClassForm(request.POST)
        if form.is_valid():
            # 增加题目
            if form.cleaned_data['problems']:
                ids = form.cleaned_data['problems'].split(',')
                try:
                    with transaction.atomic():
                        for pk in ids:
                            try:
                                course.problems.add(pk)
                            except:
                                continue
                except:
                    return HttpResponse(0)
            # 添加学生
            for stu_detail in request.POST['students'].splitlines():
                if len(stu_detail.split()) > 1:
                    id_num, username = stu_detail.split()[0], stu_detail.split()[1]
                    try:
                        student = MyUser.objects.get(id_num=id_num)
                    except:
                        student = MyUser(id_num=id_num, email=id_num + '@njupt.edu.cn', username=username)
                        student.set_password(id_num)
                        student.save()
                        student.groups.add(Group.objects.get(name='学生'))
                    # newCourse.students.add(student)
                    # IPython.embed()
                    if course.students.all().filter(codeweekclassstudent__student=student).count() == 0:
                        if course.numberEachGroup == 1:  # 没有分组操作，直接给每个学生加入一个只有自己一个人的组
                            newgroup = CodeWeekClassGroup(cwclass=course)
                            newgroup.save()
                            newCourseStudent = CodeWeekClassStudent(codeWeekClass=course, student=student,
                                                                    group=newgroup,
                                                                    isLeader=True)
                            newCourseStudent.save()
                        else:
                            newCourseStudent = CodeWeekClassStudent(codeWeekClass=course, student=student)
                            newCourseStudent.save()
                else:
                    try:
                        student = MyUser.objects.get(id_num=stu_detail)
                        if course.students.all().filter(codeweekclassstudent__student=student).count() == 0:
                            if course.numberEachGroup == 1:  # 没有分组操作，直接给每个学生加入一个只有自己一个人的组
                                newgroup = CodeWeekClassGroup(cwclass=course)
                                newgroup.save()
                                newCourseStudent = CodeWeekClassStudent(codeWeekClass=course, student=student,
                                                                        group=newgroup,
                                                                        isLeader=True)
                                newCourseStudent.save()
                            else:
                                newCourseStudent = CodeWeekClassStudent(codeWeekClass=course, student=student)
                                newCourseStudent.save()

                    except:
                        pass
            return redirect(reverse('view_course_for_teacher', args=[course.id, ]))
    else:
        form = UpdateClassForm()
        categorys = ProblemCategory.objects.all()
        return render(request, 'code_week/add_problem_student.html', {'form': form, 'categorys': categorys})

@login_required
# 用于教师获取学生的信息表格
def get_student_state(request, courseId):
    course = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
    except:
        return HttpResponse(0)
    json_data = {}
    recodes = []
    students = course.CodeWeekClass_student.all()

    json_data['total'] = students.count()

    if course.numberEachGroup == 1:
        for student in students.all():
            recode = {'pk': student.pk,
                      'name': student.get_full_name(),
                      'state': singleStuState(student),
                      'id': student.pk}
            recodes.append(recode)
    else:
        for student in students.all():
            recode = {'pk': student.pk,
                      'name': student.get_full_name(),
                      'state': multipStuState(student),
                      'id': student.pk}
            recodes.append(recode)
    json_data['rows'] = recodes
    return HttpResponse(json.dumps(json_data))

# 返回单人课程的学生状态
def singleStuState(student):
    problem = None
    try:
        problem = student.group.selectedProblem
    except:
        return "还没有选择题目"
    if not problem:
        return "还没有选择题目"
    result = "已经选择题目: " + problem.title + "<br>"
    if student.group.Code_history.all().count() == 0:
        result += "还没有提交代码"
    else:
        latest = student.group.Code_history.all()[0]
        result += "最近提交代码的时间: "
        result += latest.submitTime.strftime('%Y/%m/%d %X')
    if student.group.Report_history.all().count() == 0:
        result += "<br>" + "还没有提交课程报告"
    else:
        latest = student.group.Report_history.all()[0]
        result += "<br>" + "最近提交课程报告的时间: "
        result += latest.uploadTime.strftime('%Y/%m/%d %X')
    return result

# 返回多人课程的学生状态
def multipStuState(student):
    group = None
    try:
        group = student.group
    except:
        return "还没有加入组"
    if not group:
        return "还没有加入组"
    result = ""
    if student.isLeader:
        result += "已经建立了小组"
    else:
        for stu in group.Group_member.all():
            if stu.isLeader:
                result += '已经加入了 '
                result += stu.get_full_name()
                result += '的小组'
                break
    result += '<br>'
    result += singleStuState(student)
    return result

@login_required
# 用于教师移除学生
def remove_student(request):
    courseId = None
    studentId = None
    try:
        courseId = request.POST['courseId']
        studentId = request.POST['studentId']
    except:
        return HttpResponse(0)
    student = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
        student = CodeWeekClassStudent.objects.get(codeWeekClass=course, id=studentId)
    except:
        return HttpResponse(0)
    # 如果学生已经已经加入小组或者成立小组就无法移除
    if student.group == None:
        try:
            with transaction.atomic():
                student.delete()
        except:
            return HttpResponse(0)
        return HttpResponse(1)
    return HttpResponse(2)

@login_required
# 用于组长选择题目
def choose_problem(request, courseId):
    if request.method == 'POST':
        student = None
        problem = None
        # pdb.set_trace()
        try:
            student = CodeWeekClassStudent.objects.get(codeWeekClass=courseId, student=request.user)
            problemId = request.POST['id']
            problem = ShejiProblem.objects.get(problem_id=problemId)
        except :
            return HttpResponse(0)
        ckresult = check_time(student.codeWeekClass)
        if ckresult == TimeResult.NOTSTART:
            return HttpResponse(0)
        elif ckresult == TimeResult.FINISHED:
            return HttpResponse(0)
        try:
            with transaction.atomic():
                if student and student.group and student.isLeader and student.group.selectedProblem == None:
                    student.group.selectedProblem = problem
                    student.group.save()
                else:
                    raise Exception
        except:
            return HttpResponse(0)
        return HttpResponse(1)

@unique
class TimeResult(Enum):
    NOTSTART = 0
    FINISHED = 1
    OK = 2

def check_time(cwclass):
    # 检查现在时间是否在课程时间区间内
    nowtime = datetime.datetime.now()
    if nowtime < cwclass.begin_time:
        return TimeResult.NOTSTART
    elif nowtime > cwclass.end_time:
        return TimeResult.FINISHED
    elif cwclass.begin_time <= nowtime <= cwclass.end_time:
        return TimeResult.OK

def check_contribution(contribution, group):
    # 检查贡献度是否和为100以及贡献度中的人是否都存在
    # 这里贡献度是用,分割，格式为B14040315 zk:50,B1404031* **:50
    # pdb.set_trace()
    contributionDetails = contribution.split(',')
    if len(contributionDetails) == group.Group_member.count(): # 条目数一样
        # 开始计算和是否为100并且每个人都出现
        students_list = []
        for stu in group.Group_member.all():
            students_list.append(stu.get_full_name())
        students_list.sort()
        to_test_students_list = [] # 从贡献度字符串中得到的学生列表
        for student_contribution in contributionDetails:
            a_student_list = student_contribution.split(':')
            if not len(a_student_list) == 2:
                return False
            num = int(a_student_list[1])
            if not num == 100 and not num == 80 and not num == 60 and not num == 40 and not num == 20:
                return False
            to_test_students_list.append(a_student_list[0])
        to_test_students_list.sort()
        if students_list == to_test_students_list:
            return True
        else:
            return False
    else:
        return False

def un_zip(file_name):
    """
    解压zip文件到当前目录，并生成"file_name_files"文件夹保存解压的数据
    :param file_name: 文件名
    :return: 无
    """
    os.mkdir(file_name + '_files')
    shutil.unpack_archive(file_name, extract_dir=file_name+'_files')
    # zip_file = zipfile.ZipFile(file_name)
    # os.mkdir(file_name + "_files")
    # i = 0
    # dirname = None
    # os.chdir(file_name + "_files/")
    # #zip_file.extractall(members=zip_file.namelist())
    # for names in zip_file.namelist():
    #     #pdb.set_trace()
    #     # print(chardet.detect(names.encode('utf8')))
    #     if i == 0:
    #         dirname = names
    #         i = i + 1
    #     filename = names.encode('cp437').decode('gbk')
    #     print(filename)
    #     zip_file.extract(names)
    # #     os.chdir(file_name + "_files/")
    #     os.rename(names, filename)
    # # shutil.rmtree(file_name + "_files/" + dirname) # 有点问题，所以删掉还是有乱码的，不知道为什么
    # zip_file.close()

# 返回保存代码单文件的路径
def singleFileName(fileId):
    filename = os.path.join(USER_FILE_DIR, 'allCode', str(fileId))
    return filename

# 将解压的文件夹中的文件编号并且复制到指定目录，并且生成目录的序列化结果
def copy_generage_dict_info(work_dir, group):
    result = {}
    # pdb.set_trace()
    for root, dirs, files in os.walk(work_dir):
        testroot = root
        subdirs = []
        temp = result
        while len(work_dir) < len(testroot):  # while work_dir != dirs
            testroot, subdir = os.path.split(testroot)
            subdirs.append(subdir)
        while len(subdirs) > 0:
            sd = subdirs[-1]
            try:
                temp = temp[sd.encode('cp437').decode('gbk')]
            except UnicodeEncodeError:
                temp = temp[sd]
            subdirs.pop()
        for filename in files: # 单个文件，需要保存
            fileRecord = CodeFile.objects.create(group=group)
            shutil.copy(os.path.join(root, filename), singleFileName(fileRecord.id)) # 复制文件到指定目录
            aFileName = None
            try:
                aFileName = filename.encode('cp437').decode('gbk')
            except UnicodeEncodeError:
                aFileName = filename
            temp[aFileName] = fileRecord.id
        for dirname in dirs:
            aDirName = None
            try:
                aDirName = dirname.encode('cp437').decode('gbk')
            except UnicodeEncodeError:
                aDirName = dirname
            temp[aDirName] = {}
        # temp = temp[os.path.split(root)[1]]
        # for filename in files:
        #     temp[filename] = j
        #     ++j
        # for dirname in dirs:
        #     temp[dirname] = {}
    # print(json.dumps(result))
    return json.dumps(result)

# 返回保存代码压缩文件的路径
def codeZipFileName(fileId):
    filename = os.path.join(USER_FILE_DIR, 'codeZip', str(fileId))
    return filename

# 用于计算文件hash
def sha1(filepath):
    with open(filepath,'rb') as f:
        sha1obj = hashlib.sha1()
        sha1obj.update(f.read())
        hash = sha1obj.hexdigest()
        return hash

@login_required
# 组长提交代码，保存提交的zip文件并且将压缩包内的所有文件都解压标号，生成目录序列化字符串
def submit_code(request, courseId):
    # 检查是否是组长，现在的逻辑是只有组长才能上传代码，而且需要组长手动填写贡献量
    student = None
    try:
        student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=courseId)
    except:  # 没有查到学生或者课程
        return render(request, 'warning.html', {'info': '出现问题'})
    try:
        group = student.group
        if not group:
            return render(request, 'warning.html', {'info': '你还没有加入组或者成为组长'})
    except:
        return render(request, 'warning.html', {'info': '你还没有加入组或者成为组长'})

    if request.method == 'POST':
        # 再检查现在时间是否在课程时间之内
        ckresult = check_time(student.codeWeekClass)
        if ckresult == TimeResult.NOTSTART:
            return render(request, 'warning.html', {'info': '课程还没开始，还不能提交代码'})
        elif ckresult == TimeResult.FINISHED:
            return render(request, 'warning.html', {'info': '课程已经结束，无法提交代码'})
        elif ckresult == TimeResult.OK:  # 时间没问题
            pass
        # 检查是否是组长，现在的逻辑是只有组长才能上传代码，而且需要组长手动填写贡献量
        if not student.isLeader:
            return render(request, 'warning.html', {'info': '提交代码现在只能由组长完成'})

        form = SubmitCodeForm(request.POST, request.FILES)
        if form.is_valid():
            # 再检查贡献值是否正确，如果只是一个人无所谓
            # 如果是多个人需要检查和是否为100%
            if not check_contribution(request.POST['contribution'], student.group):
                return render(request, 'warning.html', {'info': '贡献度不正确'})
            # 处理上传的文件
            # 首先检查文件后缀是否为.zip
            # 然后开始解压压缩文件并且遍历文件夹生成文件夹的序列化结果
            # 对每个文件进行编号并且复制到保存文件夹
            # 保存序列化结果到数据库
            # 保存这个压缩文件，来方便打包文件的下载
            file = request.FILES['codeFile']
            if not file.name.endswith(".zip"):
                return render(request, 'warning.html', {'info': '只支持zip压缩文件'})
            else:
                random_name = ''.join(random.sample(string.digits + string.ascii_letters * 10, 8))
                tempdir = newProblemFileName(random_name)
                os.mkdir(tempdir)
                filename = os.path.join(tempdir, file.name)
                with open(filename, 'wb+') as destination:
                    for chunk in file.chunks():
                        destination.write(chunk)
                # un_zip(filename)
                sha1result = sha1(filename)
                # 这边可以利用文件hash做一些额外工作，例如比较和以前版本是否一样
                un_zip(filename)
                dir_result = copy_generage_dict_info(filename + '_files/', student.group)
                # 保存整个压缩文件
                newCodeZip = CodeZipFile.objects.create(fileName = file.name)
                shutil.copy(filename, codeZipFileName(newCodeZip.id))
                # 删除产生的文件
                shutil.rmtree(tempdir)
                # pdb.set_trace()
                # 保存代码版本（目录序列化结果）
                newCodeHistory = CodeDirHistory.objects.create(zipFile=newCodeZip, dirText=dir_result,
                                                               group=student.group,
                                                               contribution=request.POST['contribution'],
                                                               fileHash=sha1result)
                # 将现在小组的代码历史指向新建的
                student.group.nowCodeDir = newCodeHistory
                student.group.save()
                infoText = {'action': 'code', 'leader': student.get_full_name(), 'id': newCodeHistory.id, 'groupId': student.group.id}
                infoObject = LatestInfo.objects.create(course=student.codeWeekClass, info=json.dumps(infoText))
                updateLatestInfo(courseId, infoObject)
                return render(request, 'code_week/submit_code.html', {'members': student.group.Group_member.all(),
                                                                      'course': student.codeWeekClass})
        else:
            print(form.errors)
            return render(request, 'warning.html', {'info': '提交的信息有误'})
    else:
        if student.isLeader:
            return render(request, 'code_week/submit_code.html', {'members': student.group.Group_member.all(),
                                                              'course': student.codeWeekClass})
        else:
            return render(request, 'code_week/submit_code.html', {'course': student.codeWeekClass})

@login_required
# 返回当前组的目录结构
def get_dir_struct(request, courseId):
    student = None
    try:
        student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=courseId)
        student.group
    except:  # 没有查到学生或者课程
        return HttpResponse(json.dumps({}))
    if student.group.nowCodeDir:
        return HttpResponse(student.group.nowCodeDir.dirText)
    else:
        return HttpResponse(json.dumps({}))

def returnUtf8FileStr(fileName):
    chardetResult = get_encoding(fileName)
    print(chardetResult)
    with open(fileName, 'rb') as file:
        if chardetResult == None:
            return ""
        elif chardetResult == "utf-8" or chardetResult == 'ascii' or chardetResult == 'UTF-8-SIG':
            return file.readlines()
        else:  # 假设是gb2312
            allLine = ""
            for line in file.readlines():
                allLine += line.decode('gb2312')
            return allLine

@login_required
# 返回代码文件
def get_code_file(request, fileId):
    file = None
    try:
        file = CodeFile.objects.get(id=fileId)
    except:
        return HttpResponse("", status=404)
    # 检查请求者是否是这个文件的所有者
    for member in file.group.Group_member.all():
        if request.user == member.student:
            fileName = singleFileName(fileId)
            return HttpResponse(returnUtf8FileStr(fileName))
    return HttpResponse("")

@login_required
# 返回学生所在组的所有提交记录
def get_all_code_history(request, courseId):
    student = None
    try:
        student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=courseId)
        student.group
    except:  # 没有查到学生或者课程
        return HttpResponse(json.dumps([]))
    result = []
    for record in student.group.Code_history.all():
        result.append({'time':record.submitTime.strftime('%Y/%m/%d %X'), 'text':record.dirText,
                       'id': record.id})
    return HttpResponse(json.dumps(result))

def newReportFilename(fileId):
    filename = os.path.join(USER_FILE_DIR, 'reportFile', str(fileId))
    return filename

@login_required
# 处理组长提交课程报告
def handle_upload_report(request, courseId):
    # 检查是否是组长，现在的逻辑是只有组长才能上传代码，而且需要组长手动填写贡献量
    student = None
    try:
        student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=courseId)
    except:  # 没有查到学生或者课程
        return render(request, 'warning.html', {'info': '出现问题'})
    try:
        group = student.group
        if not group:
            return render(request, 'warning.html', {'info': '你还没有加入组或者成为组长'})
    except:
        return render(request, 'warning.html', {'info': '你还没有加入组或者成为组长'})
    ckresult = check_time(student.codeWeekClass)
    if ckresult == TimeResult.NOTSTART:
        return render(request, 'warning.html', {'info': '课程还没开始，还不能提交报告'})
    elif ckresult == TimeResult.FINISHED:
        return render(request, 'warning.html', {'info': '课程已经结束，无法提交报告'})
    elif ckresult == TimeResult.OK:  # 时间没问题
        pass
    if request.method == 'POST':
        form = UploadReportForm(request.POST, request.FILES)
        if not form.is_valid():
            return render(request, 'warning.html', {"info":"表单无效"})
        else:
            file = request.FILES['report']
            newReportFile = None
            try:
                newReportFile = ReportHistory.objects.create(filename=file.name, group=student.group)
            except:
                return render(request, 'warning.html', {"info":"保存文件时出现问题"})
            saveFilename = newReportFilename(newReportFile.id)
            with open(saveFilename, 'wb+') as destination:
                for chunk in file.chunks():
                    destination.write(chunk)
            infoText = {'action': 'report', 'leader': student.get_full_name(), 'id': newReportFile.id}
            infoObject = LatestInfo.objects.create(course=student.codeWeekClass, info=json.dumps(infoText))
            updateLatestInfo(courseId, infoObject)
            return HttpResponseRedirect(reverse('submit_code', args=(student.codeWeekClass.id,)))
    else:
        return render(request, 'warning.html', {"info":"不支持的HTTP方法"})

@login_required
# 返回学生提交的所有课程报告记录
def get_all_report_history(request, courseId):
    student = None
    try:
        student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=courseId)
        student.group
    except:  # 没有查到学生或者课程
        return HttpResponse(json.dumps([]))
    result = []
    for record in student.group.Report_history.all():
        result.append({'time':record.uploadTime.strftime('%Y/%m/%d %X'),'id': record.id})
    return HttpResponse(json.dumps(result))

@login_required
# 提供学生下载自己以前提交过的课程报告
def download_report(request, historyId):
    report = None
    try:
        report = ReportHistory.objects.get(id=historyId)
    except:
        return render(request, 'warning.html', {'info': '找不到课程报告'})
    if report:  # 能够匹配到课程报告，检查学生是否是这个课程报告的所有者
        student = None
        try:
            student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=report.group.cwclass)
        except:
            return render(request, 'warning.html', {'info': '你不是该文件的所有者'})
        owner = False
        for stu in report.group.Group_member.all():
            if student == stu:
                owner = True
                break
        if not owner:
            return render(request, 'warning.html', {'info': '你不是该文件的所有者'})

        filename = newReportFilename(report.id)

        def file_iterator(file_name, chunk_size=512):
            with open(file_name, 'rb') as f:
                while True:
                    c = f.read(chunk_size)
                    if c:
                        yield c
                    else:
                        break

        response = StreamingHttpResponse(file_iterator(filename))
        response['Content-Disposition'] = '''attachment;filename*= UTF-8''{0}'''.format(encodeFilename(report.filename))
        return response
    else:
        return render(request, 'warning.html', {'info': '找不到课程报告'})

# 实现代码打包文件的下载功能
@login_required
def download_codeZip(request, historyId):
    codeHistory = None
    try:
        codeHistory = CodeDirHistory.objects.get(id = historyId)
    except:
        return render(request, 'warning.html', {'info': '找不到代码文件'})
    if codeHistory: #能够匹配到历史代码,检查学生是否是这个代码的所有者
        student = None
        try:
            student = CodeWeekClassStudent.objects.get(student=request.user, codeWeekClass=codeHistory.group.cwclass)
        except:
            return render(request, 'warning.html', {'info': '你不是下载该文件'})
        owner = False
        for stu in codeHistory.group.Group_member.all():
            if student == stu:
                owner = True
                break
        if not owner:
            return render(request, 'warning.html', {'info': '你不是下载该文件'})

        file = codeHistory.zipFile
        filename = codeZipFileName(file.id)

        def file_iterator(file_name, chunk_size=512):
            with open(file_name, 'rb') as f:
                while True:
                    c = f.read(chunk_size)
                    if c:
                        yield c
                    else:
                        break

        response = StreamingHttpResponse(file_iterator(filename))
        response['Content-Disposition'] = '''attachment;filename*= UTF-8''{0}'''.format(encodeFilename(file.fileName))
        return response
    else:
        return render(request, 'warning.html' ,{'info' : '找不到代码文件'})

# 老师查看课程学生的代码提交情况
@login_required
def teacher_read_code(request, courseId, groupId):
    # 检查是否是该课程的老师
    course = None
    try:
        course = CodeWeekClass.objects.get(teacher=request.user, id=courseId)
    except:
        return render(request, 'warning.html', {'info': '没有查到这个课程'})
    group = None
    try:
        group = CodeWeekClassGroup.objects.get(id=groupId)
    except:
        return render(request, 'warning.html', {'info': '没有查到学生组'})
    if group.cwclass != course or group.using == False:
        return render(request, 'warning.html', {'info': '没有查到学生组'})
    leader = None
    members = []
    for stu in group.Group_member.all():
        if stu.isLeader:
            leader = stu.get_full_name()
        else:
            members.append(stu.get_full_name())
    problem = None
    if group.selectedProblem:
        problem = group.selectedProblem.title
    members.sort()
    return render(request, 'code_week/teacher_code.html', {'courseId': courseId, 'groupId': groupId,
                                                           'leader': leader, 'members': members, 'problem': problem})

# 老师查看自己课程的一个组的代码提交情况
@login_required
def teacher_get_all_code_history(request, courseId, groupId):
    # 检查是否是该课程的老师
    course = None
    try:
        course = CodeWeekClass.objects.get(teacher=request.user, id=courseId)
    except:
        return HttpResponse("")
    group = None
    try:
        group = CodeWeekClassGroup.objects.get(id=groupId)
    except:
        return HttpResponse("")
    result = []
    for record in group.Code_history.all():
        result.append({'time': record.submitTime.strftime('%Y/%m/%d %X'), 'text': record.dirText,
                       'id': record.id})
    return HttpResponse(json.dumps(result))

# 老师查看自己课程的一个组的课程报告提交情况
@login_required
def teacher_get_all_report_history(request, courseId, groupId):
    # 检查是否是该课程的老师
    course = None
    try:
        course = CodeWeekClass.objects.get(teacher=request.user, id=courseId)
    except:
        return HttpResponse("")
    group = None
    try:
        group = CodeWeekClassGroup.objects.get(id=groupId)
    except:
        return HttpResponse("")
    result = []
    for record in group.Report_history.all():
        result.append({'time': record.uploadTime.strftime('%Y/%m/%d %X'),'id': record.id})
    return HttpResponse(json.dumps(result))

# 判断文件编码
def get_encoding(file):
    with open(file, 'rb') as f:
        return chardet.detect(f.read())['encoding']

# 老师下载自己课程的某个组提交的代码
@login_required
def teacher_get_single_code(request, courseId, fileId):
    # 检查是否是该课程的老师
    course = None
    try:
        course = CodeWeekClass.objects.get(teacher=request.user, id=courseId)
    except:
        return HttpResponse("", status=404)
    file = None
    try:
        file = CodeFile.objects.get(id=fileId)
    except:
        return HttpResponse("", status=404)
    if file.group.cwclass == course:
        fileName = singleFileName(fileId)
        return HttpResponse(returnUtf8FileStr(fileName))

# 老师下载自己课程的某个组的某次提交的代码打包文件
@login_required
def teacher_get_code_zip(request, courseId, historyId):
    # 检查是否是该课程的老师
    course = None
    try:
        course = CodeWeekClass.objects.get(teacher=request.user, id=courseId)
    except:
        return HttpResponse("", status=404)
    codeHistory = None
    try:
        codeHistory = CodeDirHistory.objects.get(id=historyId)
    except:
        return render(request, 'warning.html', {'info': '找不到代码文件'})
    if codeHistory:
        if codeHistory.group.cwclass == course:
            file = codeHistory.zipFile
            filename = codeZipFileName(file.id)

            def file_iterator(file_name, chunk_size=512):
                with open(file_name, 'rb') as f:
                    while True:
                        c = f.read(chunk_size)
                        if c:
                            yield c
                        else:
                            break

            response = StreamingHttpResponse(file_iterator(filename))
            response['Content-Disposition'] = '''attachment;filename*= UTF-8''{0}'''.format(
                encodeFilename(file.fileName))
            return response
    return render(request, 'warning.html', {'info': '找不到代码文件'})

# 老师下载自己课程的某个组的某个课程报告
@login_required
def teacher_download_report(request, courseId, historyId):
    # 检查是否是该课程的老师
    course = None
    try:
        course = CodeWeekClass.objects.get(teacher=request.user, id=courseId)
    except:
        return HttpResponse("", status=404)
    report = None
    try:
        report = ReportHistory.objects.get(id=historyId)
    except:
        return render(request, 'warning.html', {'info': '找不到课程报告'})
    if report:
        if report.group.cwclass == course:
            filename = newReportFilename(report.id)

            def file_iterator(file_name, chunk_size=512):
                with open(file_name, 'rb') as f:
                    while True:
                        c = f.read(chunk_size)
                        if c:
                            yield c
                        else:
                            break

            response = StreamingHttpResponse(file_iterator(filename))
            response['Content-Disposition'] = '''attachment;filename*= UTF-8''{0}'''.format(
                encodeFilename(report.filename))
            return response
    return render(request, 'warning.html', {'info': '找不到课程报告'})

# 老师检查是否可以查看学生代码
@login_required
def teacher_check_student(request):
    courseId = None
    studentId = None
    try:
        courseId = request.POST['courseId']
        studentId = request.POST['studentId']
    except:
        return HttpResponse(-1)
    student = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
        student = CodeWeekClassStudent.objects.get(codeWeekClass=course, id=studentId)
    except:
        return HttpResponse(-1)
    # 如果学生已经已经加入小组或者成立小组就无法移除
    if student.group:
        if student.group.selectedProblem:
            return HttpResponse(student.group.id)
    return HttpResponse(-1)

# 用于教师打包所有材料
def tarFiles(courseId, className, teacherName):
    workDir = os.path.join(USER_FILE_DIR, 'codeWeekTarFiles', className + "_" + teacherName)
    if os.path.exists(workDir):
        shutil.rmtree(workDir, ignore_errors=True)
    os.mkdir(workDir)
    # os.chdir(workDir)
    studentDir = os.path.join(workDir, className + "_学生材料")
    os.mkdir(studentDir)
    os.chdir(studentDir)
    course = None
    try:
        course = CodeWeekClass.objects.get(id=courseId)
    except:
        return
    if course.numberEachGroup == 1:  # 单人模式
        for student in course.students.all().order_by("id_num"):
            cwStudent = None
            try:
                cwStudent = CodeWeekClassStudent.objects.get(codeWeekClass=course, student=student)
            except:
                return # 可以增加log异常信息
            os.mkdir(student.id_num)
            os.chdir(student.id_num)
            try:
                latestSubmit = cwStudent.group.nowCodeDir
                fileId = latestSubmit.zipFile.id
                problem = cwStudent.group.selectedProblem
                if latestSubmit and problem:
                    shutil.copy(codeZipFileName(fileId), problem.title + ".zip")
                latestReport = cwStudent.group.Report_history.all().order_by("-id")[0]
                if latestReport:
                    fileSuffix = ""
                    index = latestReport.filename.rfind('.')
                    if not index == -1:
                        fileSuffix = latestReport.filename[index:]
                    shutil.copy(newReportFilename(latestReport.id), "程序设计_" + student.id_num + fileSuffix)
            except:
                # os.chdir(BASE_DIR)
                pass # 可以增加异常信息
            os.chdir("../")
    else: # 多人模式
        groups = CodeWeekClassGroup.objects.filter(cwclass=course, using=True).order_by("id")
        groupNumber = 1
        memberNumber = 1
        f = xlwt.Workbook()
        sheet1 = f.add_sheet('sheet1', cell_overwrite_ok=True)
        sheet1.write(0, 0, '小组')
        sheet1.write(0, 1, '学号')
        sheet1.write(0, 2, '姓名')
        sheet1.write(0, 3, '贡献度比例')
        sheet1.write(0, 4, '贡献度详情')
        appearStudents = set()
        allStudents = set()
        for student in course.CodeWeekClass_student.all():
            allStudents.add(student.get_full_name())
        for group in groups:
            newMemberNumber = memberNumber + group.Group_member.count()
            sheet1.write_merge(memberNumber, newMemberNumber-1, 0, 0, groupNumber)
            groupNumber += 1
            memberLineMap = {}
            contribution_sum = {}
            for student in group.Group_member.order_by("-isLeader"):
                sheet1.write(memberNumber, 1, student.student.id_num)
                sheet1.write(memberNumber, 2, student.student.username)
                memberLineMap[student.get_full_name()] = memberNumber
                contribution_sum[student.get_full_name()] = 0
                memberNumber += 1
                appearStudents.add(student.get_full_name())
            index = 4 # 贡献度开始的列数
            for history in group.Code_history.all():
                contribution_list = history.contribution.split(',')
                for single_member_contribution in contribution_list:
                    contribution_record = single_member_contribution.split(':')
                    contribution_sum[contribution_record[0]] += int(contribution_record[1])
                    sheet1.write(memberLineMap[contribution_record[0]], index, int(contribution_record[1]))
                index += 1
            count = group.Code_history.all().count()
            for student in group.Group_member.all():
                contribution_average = 0
                if not count == 0:
                    contribution_average = contribution_sum[student.get_full_name()] / count
                sheet1.write(memberLineMap[student.get_full_name()], 3, contribution_average)
        # 记录未在分组中出现的学生
        remainStudents = allStudents - appearStudents
        if not len(remainStudents) == 0:
            sheet1.write_merge(memberNumber, memberNumber, 0, 1, "没有加入小组的学生")
            for student in allStudents - appearStudents:
                memberNumber += 1
                idAndName = student.split(' ')
                if not len(idAndName) == 2:
                    continue
                sheet1.write(memberNumber, 1, idAndName[0])
                sheet1.write(memberNumber, 2, idAndName[1])
                sheet1.write(memberNumber, 3, 0)
        f.save(os.path.join(workDir, '贡献度.xls'))
        index = 1
        for group in groups:
            NoStr = str(index)
            if len(NoStr) == 1:
                NoStr = "0"+NoStr
            dirName = "第" + NoStr + "小组_"
            if group.selectedProblem:
                dirName = dirName + group.selectedProblem.title
            else: # 小组没有选择题目
                continue
            os.mkdir(dirName)
            os.chdir(dirName)
            leader = None
            mem = []
            for member in group.Group_member.all():
                if member.isLeader:
                    leader = member
                else:
                    mem.append(member)
            mem.sort(key = lambda x:  x.student.id_num)
            reportFileName = "程序设计_"
            leaderId = leader.student.id_num[-2:]
            reportFileName += leaderId
            for member in mem:
                reportFileName += "_"
                reportFileName += member.student.id_num[-2]
            try:
                latestReport = group.Report_history.all().order_by("-id")[0]
                if latestReport:
                    fileSuffix = ""
                    index = latestReport.filename.rfind('.')
                    if not index == -1:
                        fileSuffix = latestReport.filename[index:]
                    shutil.copy(newReportFilename(latestReport.id), reportFileName + fileSuffix)
            except:
                pass
            try:
                latestSubmit = group.nowCodeDir
                fileId = latestSubmit.zipFile.id
                problem = group.selectedProblem
                if latestSubmit and problem:
                    shutil.copy(codeZipFileName(fileId), problem.title + ".zip")
            except:
                pass
            os.chdir("../")
            index += 1
    os.chdir("../")
    # 尝试合并题目word文档,python-docx只支持docx
    docxs = []
    for problem in course.problems.all():
        shutil.copy(newProblemFileName(problem.problem_id), os.path.join(workDir, problem.filename))
        docxs.append(problem.filename)
    good = True
    for d in docxs:
        try:
            doc = docx.Document(d)
        except:  # 不支持的文件
            good = False
    if good and len(docxs) >= 1:
        doc1 = docx.Document(docxs[0])
        for i in range(1, len(docxs)):
            doc2 = docx.Document(docxs[i])
            for element in doc2.element.body:
                doc1.element.body.append(element)
        doc1.save(os.path.join(workDir, "合并的文件.docx"))
    shutil.make_archive(os.path.join("..",className+"_"+teacherName), format="zip", root_dir=os.path.dirname(workDir), base_dir=className+"_"+teacherName)
    os.chdir("../")
    newTar = TarHistory.objects.create(course=course,filename=className+"_"+teacherName+".zip")
    source = os.path.join(USER_FILE_DIR, "codeWeekTarFiles", className+"_"+teacherName+".zip")
    target = os.path.join(USER_FILE_DIR, "codeWeekTarFiles", str(newTar.id))
    # pdb.set_trace()
    shutil.move(source, target)
    # shutil.copy(os.path.join(BASE_DIR,"codeWeekTarFiles",className+"_"+teacherName+".zip"), os.path.join(BASE_DIR,"codeWeekTarFiles",str(newTar.id)))
    shutil.rmtree(workDir)
    return good

# 处理教师打包请求
@login_required
def handelTeacherTar(request, courseId):
    if request.method == "POST":
        course = None
        className = None
        teacherName = None
        # pdb.set_trace()
        try:
            course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
            className = request.POST['class_name']
            teacherName = request.POST['teacher_name']
        except:
            return HttpResponse(-1)
        ckResult = check_time(course)
        if ckResult == TimeResult.NOTSTART:
            return HttpResponse(1) # 课程还没开始，不要打包了
        if tarFiles(courseId,str(className), str(teacherName)):
            return HttpResponse(0) # 打包成功，并且还合并了题目文档
        else:
            return HttpResponse(2) # 打包成功，但是无法合并题目文档

# 教师下载打包文件
@login_required
def teacherDownloadTar(request, courseId):
    course = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
    except:
        return
    historys = TarHistory.objects.filter(course=course).order_by("-id")
    if historys.count() != 0:
        latest = historys[0]
        filename = os.path.join(USER_FILE_DIR, "codeWeekTarFiles", str(latest.id))

        def file_iterator(file_name, chunk_size=512):
            with open(file_name, 'rb') as f:
                while True:
                    c = f.read(chunk_size)
                    if c:
                        yield c
                    else:
                        break

        response = StreamingHttpResponse(file_iterator(filename))
        response['Content-Disposition'] = '''attachment;filename*= UTF-8''{0}'''.format(
            encodeFilename(latest.filename))
        return response

# 教师查看课程最新信息
@login_required
def teacherViewLatestInfo(request, courseId):
    course = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
    except:
        return
    return render(request, 'code_week/latest_info.html', {'course': course})

# 教师Ajax获取课程的最新信息
@login_required
def teacherGetLatestInfo(request, courseId):
    course = None
    try:
        course = CodeWeekClass.objects.get(id=courseId, teacher=request.user)
    except:
        return
    results = []
    records = LatestInfo.objects.filter(course=course).order_by("time")
    for record in records:
        results.append({"time": record.time.strftime('%Y-%m-%d %H:%M:%S'),
                        "content": record.info})
    return HttpResponse(json.dumps(results))

# 返回所有提交的贡献度情况
@login_required
def teacherGetContribution(request, groupId):
    group = None
    try:
        group = CodeWeekClassGroup.objects.get(id=groupId)
    except:
        return HttpResponse()
    if not group.cwclass.teacher == request.user:
        return HttpResponse()
    leader = None
    members = []
    for stu in group.Group_member.all():
        if stu.isLeader:
            leader = stu.get_full_name()
        else:
            members.append(stu.get_full_name())
    members.sort()
    student_map = {}
    student_map[leader] = 1
    i = 2
    for member in members:
        student_map[member] = i
        i = i + 1
    records = []
    total = group.Code_history.all().count()
    for history in group.Code_history.all().order_by('-id'):
        record = {}
        record['time'] = history.submitTime.strftime('%Y/%m/%d %X')
        for contribution_record in history.contribution.split(','):
            name = contribution_record.split(':')[0]
            contribution = contribution_record.split(':')[1]
            record[student_map.get(name)] = contribution
        records.append(record)
    return HttpResponse(json.dumps({'rows': records, 'total': total}))

# # 通过文件夹创建dict
# import os, json
#
# work_dir = "C:\\Users\\张柯\\Desktop\\测试"
# result = {}
# i = 0
# j = 1
# for root, dirs, files in os.walk(work_dir):
#     testroot = root
#     subdirs = []
#     temp = result
#     while len(work_dir) < len(testroot):  # while work_dir != dirs
#         testroot, subdir = os.path.split(testroot)
#         subdirs.append(subdir)
#     while len(subdirs) > 0:
#         temp = temp[subdirs.pop()]
#     for filename in files:
#         temp[filename] = j
#         j = j + 1
#     for dirname in dirs:
#         temp[dirname] = {}
#     # temp = temp[os.path.split(root)[1]]
#     # for filename in files:
#     #     temp[filename] = j
#     #     ++j
#     # for dirname in dirs:
#     #     temp[dirname] = {}
# print(json.dumps(result))
#
# @login_required
# # 用来测试，返回文件夹结构
# def get_dir(request):
#     work_dir = "C:\\Users\\张柯\\Desktop\\测试"
#     result = {}
#     i = 0
#     j = 1
#     for root, dirs, files in os.walk(work_dir):
#         testroot = root
#         subdirs = []
#         temp = result
#         while len(work_dir) < len(testroot):  # while work_dir != dirs
#             testroot, subdir = os.path.split(testroot)
#             subdirs.append(subdir)
#         while len(subdirs) > 0:
#             temp = temp[subdirs.pop()]
#         for filename in files:
#             temp[filename] = j
#             j = j + 1
#         for dirname in dirs:
#             temp[dirname] = {}
#         # temp = temp[os.path.split(root)[1]]
#         # for filename in files:
#         #     temp[filename] = j
#         #     ++j
#         # for dirname in dirs:
#         #     temp[dirname] = {}
#     return HttpResponse(json.dumps(result))


